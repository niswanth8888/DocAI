import re
from pathlib import Path
from typing import TypedDict, Any
import fitz  # PyMuPDF
from docx import Document

from app.utils import clean_text


class TextUnit(TypedDict):
    page_number: int
    section_heading: str
    paragraph_index: int
    line_start: int
    line_end: int
    char_start: int
    char_end: int
    exact_text: str


class PageText(TypedDict):
    page: int
    text: str
    text_units: list[TextUnit]


SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt"}


def validate_file_extension(filename: str) -> str:
    suffix = Path(filename).suffix.lower()
    if suffix not in SUPPORTED_EXTENSIONS:
        raise ValueError(f"Unsupported file type '{suffix}'. Supported formats: PDF, DOCX, TXT.")
    return suffix


def detect_heading(text: str) -> bool:
    cleaned = text.strip()
    if not cleaned:
        return False
    # Length filter: headings are typically short
    if len(cleaned) > 100:
        return False
    # Numbered headings (e.g., "1. Introduction", "2.3 Disclaimers", "Section 4:")
    numbered_pattern = r'^(?:Section\s+\d+|[A-Z\d]+(?:\.[A-Z\d]+)*)\b'
    if re.match(numbered_pattern, cleaned, re.IGNORECASE):
        return True
    # Heuristics: Title Case or UPPERCASE, not ending with typical sentence punctuation
    if cleaned.isupper() or (cleaned[0].isupper() and not cleaned.endswith('.')):
        # Make sure it's not a sentence ending in a trailing particle or conjunction
        if not cleaned.endswith((',', ';', 'and', 'the', 'of', 'in', 'to', 'for', 'with', 'on', 'at', 'by', 'an', 'a')):
            return True
    return False


def extract_pdf(path: Path) -> list[PageText]:
    pages: list[PageText] = []
    current_heading = "General"
    
    with fitz.open(path) as doc:
        for page_index, page in enumerate(doc, start=1):
            page_text = page.get_text("text")
            blocks = page.get_text("blocks")
            text_units = []
            
            line_counter = 0
            char_offset = 0
            
            for block in blocks:
                block_text = block[4].strip()
                block_no = block[5]
                if not block_text:
                    continue
                    
                lines = block_text.split('\n')
                for line in lines:
                    line_stripped = line.strip()
                    if not line_stripped:
                        continue
                        
                    line_counter += 1
                    
                    if detect_heading(line_stripped):
                        current_heading = line_stripped
                        
                    char_len = len(line_stripped)
                    
                    unit: TextUnit = {
                        "page_number": page_index,
                        "section_heading": current_heading,
                        "paragraph_index": block_no,
                        "line_start": line_counter,
                        "line_end": line_counter,
                        "char_start": char_offset,
                        "char_end": char_offset + char_len,
                        "exact_text": line_stripped
                    }
                    text_units.append(unit)
                    char_offset += char_len + 1
                    
            clean_ptext = clean_text(page_text)
            if clean_ptext:
                pages.append({
                    "page": page_index,
                    "text": clean_ptext,
                    "text_units": text_units
                })
    return pages


def extract_docx(path: Path) -> list[PageText]:
    doc = Document(str(path))
    text_units = []
    current_heading = "General"
    current_char_offset = 0
    
    page_number = 1
    
    for idx, paragraph in enumerate(doc.paragraphs, start=1):
        text = paragraph.text.strip()
        if not text:
            continue
            
        # Heading styles or detection
        is_heading = paragraph.style.name.startswith("Heading") or detect_heading(text)
        if is_heading:
            current_heading = text
            
        char_len = len(text)
        line_count = text.count('\n') + 1
        
        # Paginate every 3000 characters
        if current_char_offset + char_len > 3000:
            page_number += 1
            current_char_offset = 0
            
        unit: TextUnit = {
            "page_number": page_number,
            "section_heading": current_heading,
            "paragraph_index": idx,
            "line_start": 1,
            "line_end": line_count,
            "char_start": current_char_offset,
            "char_end": current_char_offset + char_len,
            "exact_text": text
        }
        text_units.append(unit)
        current_char_offset += char_len + 1
        
    for t_idx, table in enumerate(doc.tables, start=1):
        for r_idx, row in enumerate(table.rows, start=1):
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if not row_text:
                continue
            
            char_len = len(row_text)
            if current_char_offset + char_len > 3000:
                page_number += 1
                current_char_offset = 0
                
            unit: TextUnit = {
                "page_number": page_number,
                "section_heading": current_heading,
                "paragraph_index": len(doc.paragraphs) + t_idx,
                "line_start": r_idx,
                "line_end": r_idx,
                "char_start": current_char_offset,
                "char_end": current_char_offset + char_len,
                "exact_text": row_text
            }
            text_units.append(unit)
            current_char_offset += char_len + 1
            
    # Group text units by page_number
    pages = []
    unique_pages = sorted(list(set(u["page_number"] for u in text_units)))
    for p in unique_pages:
        p_units = [u for u in text_units if u["page_number"] == p]
        p_text = "\n".join(u["exact_text"] for u in p_units)
        pages.append({
            "page": p,
            "text": clean_text(p_text),
            "text_units": p_units
        })
    return pages


def extract_txt(path: Path) -> list[PageText]:
    try:
        text = path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        text = path.read_text(encoding="latin-1")
        
    lines = text.splitlines()
    text_units = []
    current_heading = "General"
    current_char_offset = 0
    page_number = 1
    page_lines_count = 0
    
    for idx, line in enumerate(lines, start=1):
        line_stripped = line.strip()
        if not line_stripped:
            continue
            
        if detect_heading(line_stripped):
            current_heading = line_stripped
            
        char_len = len(line_stripped)
        
        # Paginate every 50 lines
        if page_lines_count >= 50:
            page_number += 1
            page_lines_count = 0
            
        unit: TextUnit = {
            "page_number": page_number,
            "section_heading": current_heading,
            "paragraph_index": idx,
            "line_start": idx,
            "line_end": idx,
            "char_start": current_char_offset,
            "char_end": current_char_offset + char_len,
            "exact_text": line_stripped
        }
        text_units.append(unit)
        current_char_offset += char_len + 1
        page_lines_count += 1
        
    pages = []
    unique_pages = sorted(list(set(u["page_number"] for u in text_units)))
    for p in unique_pages:
        p_units = [u for u in text_units if u["page_number"] == p]
        p_text = "\n".join(u["exact_text"] for u in p_units)
        pages.append({
            "page": p,
            "text": clean_text(p_text),
            "text_units": p_units
        })
    return pages


def extract_document_text(path: Path) -> list[PageText]:
    suffix = validate_file_extension(path.name)

    if suffix == ".pdf":
        pages = extract_pdf(path)
    elif suffix == ".docx":
        pages = extract_docx(path)
    elif suffix == ".txt":
        pages = extract_txt(path)
    else:
        pages = []

    if not pages:
        raise ValueError("No readable text found in the uploaded document.")

    return pages
